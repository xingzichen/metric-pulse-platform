"""网页、CSV、PDF、Word 和图片来源的获取与规范化。

管线先尝试受限 HTTP 下载，再按媒体类型提取正文；遇到挑战页、内容不足或特定状态码时才
启用浏览器渲染。所有下载均受大小、页数、图片数量和并发限制，并经过公共地址校验。缓存
只减少来源获取成本，不改变逐行双模型核验要求。
"""

from __future__ import annotations

import asyncio
import contextlib
import copy
import csv
import fcntl
import hashlib
import io
import json
import os
import re
import subprocess
import tempfile
import time
import zipfile
from collections.abc import Awaitable, Callable, Sequence
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any
from urllib.parse import parse_qsl, urlencode, urljoin, urlparse

import httpx
import pymupdf
import trafilatura
from bs4 import BeautifulSoup, Tag
from charset_normalizer import from_bytes
from docx import Document
from PIL import Image, ImageDraw, ImageOps

from .config import get_settings
from .forbes_ai50 import ForbesAI50Error, compact_forbes_ai50_html

MAX_DOCUMENT_BYTES = 20_000_000
MAX_DOCUMENT_CHARS = 80_000
MAX_PDF_PAGES = 50
MAX_IMAGES_PER_SOURCE = 6
MAX_VISION_IMAGES = 6
BOILERPLATE_TAGS = {
    "script",
    "style",
    "nav",
    "header",
    "footer",
    "aside",
    "form",
    "iframe",
    "noscript",
    "svg",
    "button",
}
NOISY_IMAGE_PATTERN = re.compile(
    r"(?:logo|icon|avatar|banner|advert|tracking|pixel|spacer|sprite|wechat|qrcode|qr-code)",
    re.IGNORECASE,
)
CHALLENGE_PAGE_PATTERN = re.compile(
    r"(?:verify you are human|checking your browser|attention required|access denied|"
    r"unusual traffic|security check|captcha|cf-chl-|challenge-platform|"
    r"人机验证|安全验证|请输入验证码|访问过于频繁|异常访问|访问受限)",
    re.IGNORECASE,
)
BROWSER_RETRYABLE_STATUSES = {403, 406, 408, 409, 425, 429, 500, 502, 503, 504}
BROWSER_EXCLUDED_SUFFIXES = {
    ".7z",
    ".csv",
    ".doc",
    ".docx",
    ".gif",
    ".jpeg",
    ".jpg",
    ".json",
    ".pdf",
    ".png",
    ".rar",
    ".svg",
    ".tsv",
    ".webp",
    ".xls",
    ".xlsx",
    ".xml",
    ".zip",
}
_SOURCE_CACHE: dict[str, SourceDocument] = {}
_SOURCE_CACHE_LOCKS: dict[str, asyncio.Lock] = {}
_SOURCE_CACHE_MAX_ITEMS = 2_048
_PARSER_VERSION = "source-pipeline-v5-shared-cache-images-and-cooldowns"
IMAGE_MARKER_PATTERN = re.compile(r"\[\[METRIC_PULSE_IMAGE:\d+\]\]")
_TRACKING_QUERY_KEYS = {
    "fbclid",
    "gclid",
    "mc_cid",
    "mc_eid",
    "ref",
    "ref_src",
}


@dataclass(slots=True)
class ImageReference:
    """一张网页正文图片的下载地址、邻近说明和正文占位符。"""

    url: str
    description: str
    marker: str


@dataclass(slots=True)
class ImageEvidence:
    """可送入视觉模型的一张规范化 PNG，并保留其来源编号。"""

    label: str
    png: bytes
    source_index: int
    description: str = ""
    marker: str | None = None


@dataclass(slots=True)
class SourceDocument:
    """一个候选来源经过下载、解析或浏览器渲染后的统一表示。

    ``requested_url`` 保留原始输入，``url`` 是最终跳转地址，``normalized_url`` 用于缓存键；
    三者分开保存，才能同时满足用户可追溯、内容去重和重定向审计。
    """

    index: int
    url: str
    requested_url: str | None = None
    title: str | None = None
    snippet: str | None = None
    media_type: str = "unknown"
    text: str = ""
    images: list[ImageEvidence] = field(default_factory=list)
    image_table_results: list[dict[str, Any]] = field(default_factory=list)
    error: str | None = None
    http_status: int | None = None
    retry_after_seconds: float | None = None
    source_cooldown_until: float | None = None
    source_failure_category: str | None = None
    browser_rendered: bool = False
    browser_fallback_reason: str | None = None
    cache_hit: bool = False
    persistent_cache_hit: bool = False
    normalized_url: str | None = None
    cache_key: str | None = None
    content_hash: str | None = None
    parser_version: str = _PARSER_VERSION


class BrowserChallengeError(RuntimeError):
    """公开页面要求真人验证；采集器记录失败但不会尝试绕过。"""


def normalize_source_url(url: str) -> str:
    """生成稳定获取 URL，并优先转向来源官方提供的机器可读内容。"""

    parsed = urlparse(url.strip())
    host = (parsed.hostname or "").lower()
    path = parsed.path or "/"
    if host == "github.com":
        parts = [part for part in path.split("/") if part]
        if len(parts) >= 5 and parts[2] == "blob":
            owner, repository, _, revision, *source_path = parts
            parsed = parsed._replace(
                scheme="https",
                netloc="raw.githubusercontent.com",
                path="/".join(("", owner, repository, revision, *source_path)),
                query="",
                fragment="",
            )
            host = "raw.githubusercontent.com"
            path = parsed.path
    elif host in {"data.worldbank.org", "data.worldbank.org.cn"}:
        parts = [part for part in path.split("/") if part]
        locations = dict(parse_qsl(parsed.query)).get("locations", "")
        if (
            len(parts) == 2
            and parts[0] == "indicator"
            and re.fullmatch(r"[A-Za-z0-9_.-]+", parts[1])
            and re.fullmatch(r"[A-Za-z0-9;]+", locations)
        ):
            # 数据门户首屏只显示最近值，但同页 HTML 明确链接其官方 API。直接使用中文 API
            # 可保留完整历史序列并让地区、年份和值进入可确定性匹配的表格。
            parsed = parsed._replace(
                scheme="https",
                netloc="api.worldbank.org",
                path=f"/v2/zh/country/{locations}/indicator/{parts[1]}",
                query=urlencode({"format": "json", "per_page": 20000}),
                fragment="",
            )
            host = "api.worldbank.org"
            path = parsed.path
    port = parsed.port
    netloc = host if port is None else f"{host}:{port}"
    query = urlencode(
        [
            (key, value)
            for key, value in parse_qsl(parsed.query, keep_blank_values=True)
            if not key.lower().startswith("utm_") and key.lower() not in _TRACKING_QUERY_KEYS
        ],
        doseq=True,
    )
    return parsed._replace(
        scheme=parsed.scheme.lower(),
        netloc=netloc,
        path=path,
        query=query,
        fragment="",
    ).geturl()


def _persistent_cache_path(cache_key: str) -> Path:
    digest = hashlib.sha256(cache_key.encode()).hexdigest()
    return get_settings().source_cache_root / digest[:2] / f"{digest}.json"


def _source_cache_digest(value: str) -> str:
    return hashlib.sha256(value.encode()).hexdigest()


def _source_lock_path(cache_key: str) -> Path:
    digest = _source_cache_digest(cache_key)
    return get_settings().source_cache_root / "locks" / digest[:2] / f"{digest}.lock"


def _source_failure_path(cache_key: str) -> Path:
    digest = _source_cache_digest(cache_key)
    return get_settings().source_cache_root / "failures" / digest[:2] / f"{digest}.json"


def _source_host_state_path(hostname: str) -> Path:
    digest = _source_cache_digest(hostname.lower())
    return get_settings().source_cache_root / "hosts" / digest[:2] / f"{digest}.json"


def _source_image_path(image_hash: str) -> Path:
    return get_settings().source_cache_root / "images" / image_hash[:2] / f"{image_hash}.png"


def _read_json_file(path: Path) -> dict[str, Any] | None:
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except OSError, ValueError, TypeError:
        return None
    return payload if isinstance(payload, dict) else None


def _write_json_file(path: Path, payload: dict[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    fd, temporary_name = tempfile.mkstemp(prefix=f".{path.name}.", dir=path.parent)
    try:
        with os.fdopen(fd, "w", encoding="utf-8") as handle:
            json.dump(payload, handle, ensure_ascii=False, separators=(",", ":"))
        os.replace(temporary_name, path)
    finally:
        with contextlib.suppress(FileNotFoundError):
            os.unlink(temporary_name)


@contextlib.asynccontextmanager
async def _cross_process_source_lock(cache_key: str):
    """用共享缓存卷上的文件锁合并跨进程相同 URL 获取。"""

    path = _source_lock_path(cache_key)
    path.parent.mkdir(parents=True, exist_ok=True)
    handle = path.open("a+b")
    try:
        while True:
            try:
                fcntl.flock(handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                break
            except BlockingIOError:
                await asyncio.sleep(0.1)
        yield
    finally:
        fcntl.flock(handle.fileno(), fcntl.LOCK_UN)
        handle.close()


def _active_source_cooldown(cache_key: str, normalized_url: str) -> dict[str, Any] | None:
    """返回 URL 或域名仍生效的冷却状态，过期状态不阻止新请求。"""

    hostname = (urlparse(normalized_url).hostname or "").lower()
    states = [
        state
        for state in (
            _read_json_file(_source_failure_path(cache_key)),
            _read_json_file(_source_host_state_path(hostname)) if hostname else None,
        )
        if state and isinstance(state.get("blocked_until"), int | float)
    ]
    active = [state for state in states if float(state["blocked_until"]) > time.time()]
    return max(active, key=lambda state: float(state["blocked_until"])) if active else None


def _failure_category(document: SourceDocument) -> str | None:
    error = document.error or ""
    if document.http_status == 429:
        return "THROTTLED"
    if document.http_status == 403 or looks_like_challenge_page(error):
        return "CHALLENGE"
    if document.http_status in BROWSER_RETRYABLE_STATUSES or (
        error and "unsupported content type" not in error and "HTTP 4" not in error
    ):
        return "TRANSIENT"
    return None


def _record_source_failure(
    cache_key: str,
    normalized_url: str,
    document: SourceDocument,
) -> None:
    """把挑战/限流/瞬时错误写入共享负缓存，避免相同行继续撞站点。"""

    category = _failure_category(document)
    if category is None:
        return
    settings = get_settings()
    path = _source_failure_path(cache_key)
    previous = _read_json_file(path) or {}
    failure_count = int(previous.get("failure_count") or 0) + 1
    if category == "CHALLENGE":
        cooldown = settings.source_challenge_cooldown_seconds
    else:
        cooldown = settings.source_transient_cooldown_base_seconds * (2 ** (failure_count - 1))
    cooldown = min(settings.source_cooldown_max_seconds, max(cooldown, document.retry_after_seconds or 0))
    blocked_until = time.time() + cooldown
    state = {
        "normalized_url": normalized_url,
        "hostname": (urlparse(normalized_url).hostname or "").lower(),
        "category": category,
        "failure_count": failure_count,
        "blocked_until": blocked_until,
        "http_status": document.http_status,
        "error": (document.error or "")[:500],
        "updated_at": time.time(),
    }
    _write_json_file(path, state)
    if category in {"CHALLENGE", "THROTTLED"} and state["hostname"]:
        host_path = _source_host_state_path(state["hostname"])
        existing = _read_json_file(host_path) or {}
        if float(existing.get("blocked_until") or 0) < blocked_until:
            _write_json_file(host_path, state)
    document.source_cooldown_until = blocked_until
    document.source_failure_category = category


def _clear_source_failure(cache_key: str, normalized_url: str) -> None:
    with contextlib.suppress(FileNotFoundError):
        _source_failure_path(cache_key).unlink()
    hostname = (urlparse(normalized_url).hostname or "").lower()
    if not hostname:
        return
    host_path = _source_host_state_path(hostname)
    host_state = _read_json_file(host_path) or {}
    # 同一 URL 的 Playwright 回退成功说明先前挑战已被正常公开页面化解；清除它设置的
    # 域级冷却。其他 URL 更新的更晚冷却不能被这次成功误删。
    if host_state.get("normalized_url") == normalized_url:
        with contextlib.suppress(FileNotFoundError):
            host_path.unlink()


def _cooldown_document(
    candidate: Any,
    index: int,
    normalized_url: str,
    cache_key: str,
    state: dict[str, Any],
) -> SourceDocument:
    remaining = max(0.0, float(state["blocked_until"]) - time.time())
    return SourceDocument(
        index=index,
        url=normalized_url,
        requested_url=candidate.source_url,
        normalized_url=normalized_url,
        cache_key=cache_key,
        title=candidate.title,
        snippet=candidate.excerpt,
        error=f"source cooldown active: {state.get('category') or 'TRANSIENT'}",
        retry_after_seconds=remaining,
        source_cooldown_until=float(state["blocked_until"]),
        source_failure_category=str(state.get("category") or "TRANSIENT"),
    )


async def _reserve_host_request_slot(normalized_url: str) -> None:
    """跨进程预留域名访问时隙，使不同 URL 也遵守最小间隔。"""

    interval = get_settings().source_host_min_interval_seconds
    hostname = (urlparse(normalized_url).hostname or "").lower()
    if not hostname or interval <= 0:
        return
    path = _source_host_state_path(hostname).with_suffix(".rate.json")
    lock_path = path.with_suffix(".lock")
    lock_path.parent.mkdir(parents=True, exist_ok=True)
    handle = lock_path.open("a+b")
    try:
        while True:
            try:
                fcntl.flock(handle.fileno(), fcntl.LOCK_EX | fcntl.LOCK_NB)
                break
            except BlockingIOError:
                await asyncio.sleep(0.05)
        now = time.time()
        previous = _read_json_file(path) or {}
        scheduled_at = max(now, float(previous.get("reserved_at") or 0) + interval)
        _write_json_file(path, {"hostname": hostname, "reserved_at": scheduled_at})
    finally:
        fcntl.flock(handle.fileno(), fcntl.LOCK_UN)
        handle.close()
    if scheduled_at > now:
        await asyncio.sleep(scheduled_at - now)


def _load_persistent_document(
    cache_key: str,
    candidate: Any,
    index: int,
    normalized_url: str,
) -> SourceDocument | None:
    """读取未过期且解析器版本一致的跨任务缓存；损坏缓存按未命中处理。"""

    path = _persistent_cache_path(cache_key)
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except OSError, ValueError, TypeError:
        return None
    persisted_key = payload.get("cache_key", payload.get("normalized_url"))
    if payload.get("parser_version") != _PARSER_VERSION or persisted_key != cache_key:
        return None
    cached_at = payload.get("cached_at")
    if not isinstance(cached_at, int | float) or (
        time.time() - cached_at > get_settings().source_cache_ttl_seconds
    ):
        return None
    images: list[ImageEvidence] = []
    persisted_images = payload.get("images") if isinstance(payload.get("images"), list) else []
    for item in persisted_images:
        if not isinstance(item, dict) or not isinstance(item.get("image_hash"), str):
            continue
        try:
            png = _source_image_path(item["image_hash"]).read_bytes()
        except OSError:
            continue
        if hashlib.sha256(png).hexdigest() != item["image_hash"]:
            continue
        images.append(
            ImageEvidence(
                label=str(item.get("label") or f"Source {index}, cached image"),
                png=png,
                source_index=index,
                description=str(item.get("description") or ""),
                marker=item.get("marker") if isinstance(item.get("marker"), str) else None,
            )
        )
    image_table_results = (
        payload.get("image_table_results") if isinstance(payload.get("image_table_results"), list) else []
    )
    # 待识别图片元数据存在但对象丢失时不能返回只剩占位符的正文；重新抓取才能恢复证据。
    if persisted_images and not images and not image_table_results:
        return None
    document = SourceDocument(
        index=index,
        url=payload.get("final_url") or normalized_url,
        requested_url=candidate.source_url,
        title=payload.get("title") or candidate.title,
        snippet=payload.get("snippet") or candidate.excerpt,
        media_type=payload.get("media_type") or "unknown",
        text=payload.get("text") or "",
        images=images,
        image_table_results=image_table_results,
        http_status=payload.get("http_status"),
        browser_rendered=payload.get("browser_rendered") is True,
        browser_fallback_reason=payload.get("browser_fallback_reason"),
        cache_hit=True,
        persistent_cache_hit=True,
        normalized_url=payload.get("normalized_url") or normalized_url,
        cache_key=cache_key,
        content_hash=payload.get("content_hash"),
    )
    return document if document.text else None


def _persist_document(cache_key: str, document: SourceDocument) -> None:
    """原子写入可复用文本缓存，失败或空内容永不缓存。"""

    if document.error or not document.text:
        return
    path = _persistent_cache_path(cache_key)
    path.parent.mkdir(parents=True, exist_ok=True)
    persisted_images: list[dict[str, Any]] = []
    for image in document.images:
        image_hash = hashlib.sha256(image.png).hexdigest()
        image_path = _source_image_path(image_hash)
        if not image_path.exists():
            image_path.parent.mkdir(parents=True, exist_ok=True)
            fd, temporary_name = tempfile.mkstemp(prefix=f".{image_path.name}.", dir=image_path.parent)
            try:
                with os.fdopen(fd, "wb") as handle:
                    handle.write(image.png)
                os.replace(temporary_name, image_path)
            finally:
                with contextlib.suppress(FileNotFoundError):
                    os.unlink(temporary_name)
        persisted_images.append(
            {
                "image_hash": image_hash,
                "label": image.label,
                "description": image.description,
                "marker": image.marker,
            }
        )
    payload = {
        "parser_version": _PARSER_VERSION,
        "cached_at": time.time(),
        "cache_key": cache_key,
        "normalized_url": document.normalized_url,
        "final_url": document.url,
        "title": document.title,
        "snippet": document.snippet,
        "media_type": document.media_type,
        "text": document.text,
        "images": persisted_images,
        "image_table_results": document.image_table_results,
        "http_status": document.http_status,
        "browser_rendered": document.browser_rendered,
        "browser_fallback_reason": document.browser_fallback_reason,
        "content_hash": document.content_hash or hashlib.sha256(document.text.encode()).hexdigest(),
    }
    _write_json_file(path, payload)


def _cached_document(document: SourceDocument, candidate: Any, index: int) -> SourceDocument:
    cloned = copy.deepcopy(document)
    cloned.index = index
    cloned.requested_url = candidate.source_url
    # 来源证据元数据属于共享前缀；缓存命中时保留首次规范化结果，避免行级候选标题
    # 改变 OMLX 前缀并降低相同 URL 的 KV cache 命中率。
    cloned.title = cloned.title or candidate.title
    cloned.snippet = cloned.snippet or candidate.excerpt
    cloned.cache_hit = True
    cloned.normalized_url = normalize_source_url(candidate.source_url)
    for image in cloned.images:
        image.source_index = index
    return cloned


def persist_enriched_source_document(document: SourceDocument) -> None:
    """在图片表格化后更新进程内和跨任务缓存。

    图片模型输出已经嵌回 ``document.text``，持久化后同一来源的其他行可以
    直接复用，不再重复下载或识图。
    """

    cache_key = document.cache_key or document.normalized_url
    if not cache_key or document.error or not document.text:
        return
    if len(_SOURCE_CACHE) >= _SOURCE_CACHE_MAX_ITEMS and cache_key not in _SOURCE_CACHE:
        _SOURCE_CACHE.pop(next(iter(_SOURCE_CACHE)))
    _SOURCE_CACHE[cache_key] = copy.deepcopy(document)
    _persist_document(cache_key, document)


def _decode_text(data: bytes) -> str:
    match = from_bytes(data).best()
    return str(match) if match is not None else data.decode("utf-8", errors="replace")


def compact_json_records(data: bytes) -> str | None:
    """把 JSON 中最大的同构对象数组转成 CSV，供通用结构化匹配和模型复核。"""

    try:
        payload = json.loads(_decode_text(data))
    except json.JSONDecodeError, UnicodeError:
        return None

    candidates: list[list[dict[str, Any]]] = []

    def visit(value: Any, depth: int = 0) -> None:
        if depth > 5:
            return
        if isinstance(value, list):
            if value and all(isinstance(item, dict) for item in value):
                candidates.append(value)
            for item in value:
                visit(item, depth + 1)
        elif isinstance(value, dict):
            for item in value.values():
                visit(item, depth + 1)

    visit(payload)
    if not candidates:
        return None
    records = max(candidates, key=len)[:20_000]

    flattened: list[dict[str, Any]] = []
    columns: list[str] = []
    for record in records:
        row: dict[str, Any] = {}
        for key, value in record.items():
            if value is None or isinstance(value, str | int | float | bool):
                row[str(key)] = value
            elif isinstance(value, dict):
                # 常见统计 API 用 {id, value} 表示国家、指标等实体；把展示值放在原列名，
                # 同时保留其 id，既方便中文地区匹配也不丢失机器标识。
                if value.get("value") is None or isinstance(value.get("value"), str | int | float | bool):
                    row[str(key)] = value.get("value")
                if value.get("id") is None or isinstance(value.get("id"), str | int | float | bool):
                    row[f"{key}_id"] = value.get("id")
        for key in row:
            if key not in columns and len(columns) < 32:
                columns.append(key)
        flattened.append(row)
    if not columns:
        return None

    output = io.StringIO()
    writer = csv.writer(output, lineterminator="\n")
    writer.writerow(columns)
    for row in flattened:
        writer.writerow([row.get(column) for column in columns])
    return _normalize_text(output.getvalue())


def _normalize_text(text: str) -> str:
    lines = [re.sub(r"[ \f\v]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)[:MAX_DOCUMENT_CHARS]


def looks_like_challenge_page(text: str) -> bool:
    sample = text[:20_000]
    strong_markup = "cf-chl-" in sample.lower() or "challenge-platform" in sample.lower()
    return strong_markup or (len(text.strip()) < 6_000 and bool(CHALLENGE_PAGE_PATTERN.search(sample)))


def browser_fallback_reason(document: SourceDocument, *, min_content_chars: int) -> str | None:
    """判断是否需要浏览器重试，并返回可审计原因；附件类型不走浏览器。"""
    if document.error and document.error.startswith("source cooldown active:"):
        return None
    if (urlparse(document.url).hostname or "").lower() == "api.github.com":
        # 结构化 API 的限流或鉴权错误必须由 HTTP 重试处理；浏览器渲染既不能修复响应，也会
        # 把错误页误当证据并显著拖慢固定榜单任务。
        return None
    suffix = Path(urlparse(document.url).path).suffix.lower()
    if suffix in BROWSER_EXCLUDED_SUFFIXES:
        return None
    if document.http_status in BROWSER_RETRYABLE_STATUSES:
        return f"HTTP {document.http_status}"
    if document.media_type in {"text/html", "application/xhtml+xml"}:
        if looks_like_challenge_page(document.text):
            return "challenge-like HTML"
        if len(document.text.strip()) < min_content_chars:
            return f"main content shorter than {min_content_chars} characters"
    if document.error and document.media_type in {"unknown", "text/html", "application/xhtml+xml"}:
        return "HTTP transport failed"
    return None


def _compact_image_description(image: Tag) -> str:
    """提取带类型标签的图注、图片属性和邻近正文，供视觉模型联合判断。"""

    candidates: list[tuple[str, str]] = []
    for attribute in ("alt", "title", "data-caption", "aria-label"):
        value = image.get(attribute)
        if isinstance(value, str) and value.strip():
            candidates.append((attribute, value.strip()))
    figure = image.find_parent("figure")
    if isinstance(figure, Tag):
        caption = figure.find("figcaption")
        if isinstance(caption, Tag):
            candidates.insert(0, ("图注", caption.get_text(" ", strip=True)))

    nearby: list[str] = []
    for tag in image.find_all_previous(["p", "figcaption"], limit=3):
        text = tag.get_text(" ", strip=True)
        if text:
            nearby.append(text)
    # “图说/图表/排名”类说明比普通前文更可能真正描述当前图片。
    nearby.sort(
        key=lambda value: (
            bool(re.search(r"(?:图说|图表|排名|得分|chart|table|figure)", value, re.I)),
            -len(value),
        ),
        reverse=True,
    )
    candidates.extend(("邻近正文", value) for value in nearby[:1])
    cleaned: list[str] = []
    seen_values: set[str] = set()
    for label, value in candidates:
        normalized = re.sub(r"\s+", " ", value).strip()
        if normalized and normalized not in seen_values:
            seen_values.add(normalized)
            cleaned.append(f"{label}: {normalized[:300]}")
    return " | ".join(cleaned)[:600]


def extract_html_main_content(data: bytes, base_url: str) -> tuple[str, list[ImageReference]]:
    """抽取网页主正文，并保留最多六张候选正文图片在文本中的位置。"""

    html_text = _decode_text(data)
    # 福布斯列表页的完整 50 条数据位于 Next.js 内嵌状态而非首屏正文。先压缩为仅含业务字段
    # 的官方快照，既避免丢失未首屏渲染的公司，也不会把几十万字符页面状态送入模型。
    try:
        forbes_snapshot = compact_forbes_ai50_html(html_text)
    except ForbesAI50Error:
        # 识别到 AI 50 页面却无法解析时保留空正文；调用方会失败关闭并重试，不能退回只含
        # 少数首屏卡片的内容后误认为完整榜单。
        forbes_snapshot = ""
    if forbes_snapshot is not None:
        return _normalize_text(forbes_snapshot), []
    soup = BeautifulSoup(html_text, "lxml")
    for tag in soup.find_all(BOILERPLATE_TAGS):
        tag.decompose()
    container = soup.find("main") or soup.find("article") or soup.find(attrs={"role": "main"})
    if not isinstance(container, Tag):
        container = soup.body if isinstance(soup.body, Tag) else soup
    ranked_images: list[tuple[int, int, Tag, str, str]] = []
    for order, image in enumerate(container.find_all("img")):
        src = image.get("data-src") or image.get("src")
        if not isinstance(src, str) or not src.strip():
            continue
        description = _compact_image_description(image)
        class_text = " ".join(str(value) for value in image.get("class", []))
        absolute = urljoin(base_url, src.strip())
        if urlparse(absolute).scheme not in {"http", "https"}:
            continue
        noisy = bool(NOISY_IMAGE_PATTERN.search(absolute + " " + description + " " + class_text))
        width = int(image.get("width", 0)) if str(image.get("width", "")).isdigit() else 0
        height = int(image.get("height", 0)) if str(image.get("height", "")).isdigit() else 0
        if noisy or (width and width < 160) or (height and height < 100):
            continue
        data_hint = bool(
            re.search(r"(?:图说|图表|排名|得分|指数|chart|table|figure|rank)", description, re.I)
        )
        score = (
            (3 if data_hint else 0) + (2 if description else 0) + (1 if width >= 500 or height >= 300 else 0)
        )
        ranked_images.append((score, order, image, absolute, description))

    selected = sorted(
        sorted(ranked_images, key=lambda item: (item[0], -item[1]), reverse=True)[:MAX_IMAGES_PER_SOURCE],
        key=lambda item: item[1],
    )
    references: list[ImageReference] = []
    for position, (_, _, image, absolute, description) in enumerate(selected, start=1):
        marker = f"[[METRIC_PULSE_IMAGE:{position}]]"
        marker_tag = soup.new_tag("p")
        marker_tag.string = marker
        image.insert_after(marker_tag)
        references.append(ImageReference(url=absolute, description=description, marker=marker))

    extracted = trafilatura.extract(
        str(soup),
        include_comments=False,
        include_tables=True,
        favor_precision=True,
        no_fallback=False,
    )
    fallback = container.get_text("\n", strip=True)
    text = _normalize_text(extracted or fallback)
    return text, references


def extract_pdf(data: bytes, source_index: int) -> tuple[str, list[ImageEvidence]]:
    """提取限定页数的 PDF 文本，并渲染数字密度最高的代表页供视觉复核。"""

    document = pymupdf.open(stream=data, filetype="pdf")
    page_text: list[tuple[int, str]] = []
    for page_index in range(min(len(document), MAX_PDF_PAGES)):
        page_text.append((page_index, document[page_index].get_text("text")))
    text = _normalize_text("\n".join(value for _, value in page_text))

    # Render representative pages so charts and scanned text can participate in the same model call.
    ranked_pages = sorted(
        page_text,
        key=lambda item: (len(re.findall(r"\d", item[1])), len(item[1])),
        reverse=True,
    )
    if not ranked_pages and len(document):
        ranked_pages = [(0, "")]
    images: list[ImageEvidence] = []
    for page_index, _ in ranked_pages[:MAX_IMAGES_PER_SOURCE]:
        pixmap = document[page_index].get_pixmap(matrix=pymupdf.Matrix(1.4, 1.4), alpha=False)
        images.append(
            ImageEvidence(
                label=f"Source {source_index}, PDF page {page_index + 1}",
                png=pixmap.tobytes("png"),
                source_index=source_index,
            )
        )
    document.close()
    return text, images


def extract_docx(data: bytes, source_index: int) -> tuple[str, list[ImageEvidence]]:
    """提取 DOCX 段落、表格与有限数量的内嵌图片。"""

    document = Document(io.BytesIO(data))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            lines.append("\t".join(cell.text.strip() for cell in row.cells))
    images: list[ImageEvidence] = []
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        media_names = [name for name in archive.namelist() if name.startswith("word/media/")]
        for media_name in media_names[:MAX_IMAGES_PER_SOURCE]:
            converted = normalize_image(archive.read(media_name))
            if converted:
                images.append(
                    ImageEvidence(
                        label=f"Source {source_index}, Word image {Path(media_name).name}",
                        png=converted,
                        source_index=source_index,
                    )
                )
    return _normalize_text("\n".join(lines)), images


def extract_legacy_doc(data: bytes) -> str:
    """通过系统可用的 antiword/textutil 尽力读取旧版 DOC；失败返回空文本。"""

    with tempfile.NamedTemporaryFile(suffix=".doc") as source:
        source.write(data)
        source.flush()
        commands = [
            ["antiword", source.name],
            ["textutil", "-convert", "txt", "-stdout", source.name],
        ]
        for command in commands:
            try:
                completed = subprocess.run(command, capture_output=True, timeout=20, check=False)
            except FileNotFoundError, subprocess.TimeoutExpired:
                continue
            if completed.returncode == 0 and completed.stdout:
                return _normalize_text(_decode_text(completed.stdout))
    return ""


def normalize_image(data: bytes) -> bytes | None:
    """过滤小图，修正 EXIF 方向并缩放为模型可接受的 RGB PNG。"""

    try:
        with Image.open(io.BytesIO(data)) as image:
            image.load()
            if image.width < 160 or image.height < 100:
                return None
            converted = ImageOps.exif_transpose(image).convert("RGB")
            converted.thumbnail((1400, 1400))
            output = io.BytesIO()
            converted.save(output, format="PNG", optimize=True)
            return output.getvalue()
    except OSError, ValueError:
        return None


async def _download(
    client: httpx.AsyncClient,
    url: str,
    validate_url: Callable[[str], Awaitable[None]],
    *,
    max_bytes: int = MAX_DOCUMENT_BYTES,
) -> tuple[str, str, bytes]:
    """流式下载受限大小的公共资源，并再次校验重定向后的最终地址。"""

    await validate_url(url)
    headers = {"User-Agent": "MetricPulse/1.0"}
    if (urlparse(url).hostname or "").lower() == "api.github.com":
        headers.update(
            {
                "Accept": "application/vnd.github+json",
                "X-GitHub-Api-Version": "2022-11-28",
            }
        )
        if token := get_settings().github_api_token.strip():
            headers["Authorization"] = f"Bearer {token}"
    async with client.stream("GET", url, headers=headers) as response:
        response.raise_for_status()
        final_url = str(response.url)
        await validate_url(final_url)
        declared = int(response.headers.get("content-length", 0) or 0)
        if declared > max_bytes:
            raise ValueError(f"source exceeds {max_bytes} bytes")
        content = bytearray()
        async for chunk in response.aiter_bytes():
            content.extend(chunk)
            if len(content) > max_bytes:
                raise ValueError(f"source exceeds {max_bytes} bytes")
    media_type = response.headers.get("content-type", "").split(";", 1)[0].lower()
    return final_url, media_type, bytes(content)


async def _download_html_images(
    client: httpx.AsyncClient,
    references: list[ImageReference],
    source_index: int,
    validate_url: Callable[[str], Awaitable[None]],
) -> list[ImageEvidence]:
    images: list[ImageEvidence] = []
    for reference in references:
        try:
            _, media_type, data = await _download(
                client,
                reference.url,
                validate_url,
                max_bytes=6_000_000,
            )
        except httpx.HTTPError, ValueError:
            continue
        if not media_type.startswith("image/"):
            continue
        converted = normalize_image(data)
        if converted:
            images.append(
                ImageEvidence(
                    label=f"Source {source_index}, {reference.description or 'web image'}",
                    png=converted,
                    source_index=source_index,
                    description=reference.description,
                    marker=reference.marker,
                )
            )
    return images


async def fetch_source_document(
    candidate: Any,
    index: int,
    client: httpx.AsyncClient,
    validate_url: Callable[[str], Awaitable[None]],
) -> SourceDocument:
    """按内容类型选择解析器，把所有可预期抓取失败收敛到 document.error。

    调用者可以据此决定浏览器降级或搜索降级，而无需用异常区分 HTTP、格式和解析错误。
    """

    document = SourceDocument(
        index=index,
        url=candidate.source_url,
        requested_url=candidate.source_url,
        title=candidate.title,
        snippet=candidate.excerpt,
        normalized_url=normalize_source_url(candidate.source_url),
    )
    try:
        final_url, media_type, data = await _download(client, document.url, validate_url)
        document.url = final_url
        document.media_type = media_type or "application/octet-stream"
        document.content_hash = hashlib.sha256(data).hexdigest()
        suffix = Path(urlparse(final_url).path).suffix.lower()
        if media_type in {"text/html", "application/xhtml+xml"} or suffix in {".html", ".htm", ""}:
            document.text, references = extract_html_main_content(data, final_url)
            document.images = await _download_html_images(client, references, index, validate_url)
            if looks_like_challenge_page(f"{final_url}\n{document.text}"):
                document.url = document.requested_url or document.url
                document.text = ""
                document.images = []
                document.error = "HTTP response is an access-challenge page"
        elif media_type == "application/pdf" or suffix == ".pdf" or data.startswith(b"%PDF"):
            document.text, document.images = extract_pdf(data, index)
        elif (
            media_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or suffix == ".docx"
        ):
            document.text, document.images = extract_docx(data, index)
        elif media_type == "application/msword" or suffix == ".doc":
            document.text = extract_legacy_doc(data)
        elif media_type.startswith("image/"):
            converted = normalize_image(data)
            if converted:
                document.images = [ImageEvidence(f"Source {index}, image", converted, index)]
        elif media_type == "application/json" or media_type.endswith("+json") or suffix == ".json":
            raw_text = _normalize_text(_decode_text(data))
            # GitHub 月榜有专用完整性、排序和 Top10 解析器，必须保留原始响应；其他普通
            # JSON 来源优先压成表格，无法识别同构记录数组时再退回原文。
            document.text = (
                raw_text
                if (urlparse(final_url).hostname or "").lower() == "api.github.com"
                else compact_json_records(data) or raw_text
            )
        elif media_type.startswith("text/") or suffix in {".csv", ".tsv", ".xml"}:
            document.text = _normalize_text(_decode_text(data))
        else:
            document.error = f"unsupported content type: {media_type or suffix or 'unknown'}"
    except httpx.HTTPStatusError as exc:
        document.http_status = exc.response.status_code
        retry_after = exc.response.headers.get("retry-after", "").strip()
        if retry_after.isdigit():
            document.retry_after_seconds = min(float(retry_after), 600)
        document.error = f"HTTP {exc.response.status_code}: {exc.response.reason_phrase}"
    except (httpx.HTTPError, ValueError, OSError, RuntimeError, zipfile.BadZipFile) as exc:
        document.error = str(exc)
    return document


async def _render_document_in_browser(
    document: SourceDocument,
    context: Any,
    client: httpx.AsyncClient,
    validate_url: Callable[[str], Awaitable[None]],
    *,
    timeout_seconds: float,
    settle_seconds: float,
) -> None:
    """在隔离浏览器上下文中渲染单页，不处理真人验证挑战。

    导航请求仍经过 URL 安全校验；字体、媒体和 WebSocket 被阻止以降低成本。正文抽取后只
    截取主视口作为补充视觉证据，不把整页广告和导航截图送入模型。
    """

    page = await context.new_page()
    timeout_ms = max(1_000, int(timeout_seconds * 1_000))
    page.set_default_timeout(timeout_ms)
    try:
        response = await page.goto(
            document.requested_url or document.url,
            wait_until="domcontentloaded",
            timeout=timeout_ms,
        )
        with contextlib.suppress(Exception):  # optional on pages with long-lived requests
            await page.wait_for_load_state("networkidle", timeout=min(timeout_ms, 15_000))
        if settle_seconds > 0:
            await page.wait_for_timeout(settle_seconds * 1_000)

        final_url = page.url
        await validate_url(final_url)
        html_text = await page.content()
        text, references = extract_html_main_content(html_text.encode("utf-8"), final_url)
        title = await page.title()
        if looks_like_challenge_page(f"{final_url}\n{title}\n{text}") or looks_like_challenge_page(html_text):
            raise BrowserChallengeError("human-verification or access-challenge page detected")

        document.url = final_url
        document.title = document.title or title or None
        document.media_type = "text/html"
        document.text = text
        document.http_status = response.status if response is not None else None
        document.images = await _download_html_images(
            client,
            references,
            document.index,
            validate_url,
        )

        main = page.locator("main, article, [role='main']").first
        if await main.count():
            with contextlib.suppress(Exception):
                await main.scroll_into_view_if_needed(timeout=min(timeout_ms, 10_000))
        try:
            screenshot = await page.screenshot(
                full_page=False,
                animations="disabled",
                timeout=min(timeout_ms, 30_000),
            )
            converted = normalize_image(screenshot)
            if converted:
                document.images.insert(
                    0,
                    ImageEvidence(
                        label=f"Source {document.index}, browser-rendered main viewport",
                        png=converted,
                        source_index=document.index,
                    ),
                )
                document.images = document.images[:MAX_IMAGES_PER_SOURCE]
        except Exception:
            pass

        if not document.text and not document.images:
            raise RuntimeError("browser rendered no usable main content")
        document.browser_rendered = True
        document.error = None
    finally:
        await page.close()


async def apply_browser_fallbacks(
    documents: Sequence[SourceDocument],
    client: httpx.AsyncClient,
    validate_url: Callable[[str], Awaitable[None]],
    *,
    timeout_seconds: float,
    settle_seconds: float,
    min_content_chars: int,
    site_cooldown_seconds: float,
) -> None:
    """串行处理确有必要的浏览器降级，并按站点实施冷却时间。

    浏览器初始化失败不会终止整行采集，而是把同一错误写回所有待处理文档，供后续搜索或
    人工审核判断。
    """

    pending: list[SourceDocument] = []
    for document in documents:
        reason = browser_fallback_reason(document, min_content_chars=min_content_chars)
        if reason:
            document.browser_fallback_reason = reason
            pending.append(document)
    if not pending:
        return

    from playwright.async_api import async_playwright

    last_visit_started: dict[str, float] = {}
    try:
        async with async_playwright() as playwright:
            browser = await playwright.chromium.launch(headless=True)
            chromium_version = browser.version
            user_agent = (
                "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                f"(KHTML, like Gecko) Chrome/{chromium_version} Safari/537.36"
            )
            context = await browser.new_context(
                user_agent=user_agent,
                locale="zh-CN",
                timezone_id="Asia/Shanghai",
                viewport={"width": 1440, "height": 1000},
                color_scheme="light",
                service_workers="block",
                extra_http_headers={"Accept-Language": "zh-CN,zh;q=0.9,en;q=0.7"},
            )

            async def route_request(route: Any) -> None:
                request = route.request
                if request.resource_type in {"font", "media", "websocket"}:
                    await route.abort()
                    return
                if request.is_navigation_request():
                    try:
                        await validate_url(request.url)
                    except ValueError:
                        await route.abort()
                        return
                await route.continue_()

            await context.route("**/*", route_request)
            try:
                for document in pending:
                    hostname = urlparse(document.url).hostname or ""
                    earliest = last_visit_started.get(hostname, 0) + site_cooldown_seconds
                    delay = max(0.0, earliest - time.monotonic())
                    delay = max(delay, document.retry_after_seconds or 0)
                    if delay:
                        await asyncio.sleep(delay)
                    last_visit_started[hostname] = time.monotonic()
                    try:
                        await _render_document_in_browser(
                            document,
                            context,
                            client,
                            validate_url,
                            timeout_seconds=timeout_seconds,
                            settle_seconds=settle_seconds,
                        )
                    except BrowserChallengeError as exc:
                        document.url = document.requested_url or document.url
                        document.text = ""
                        document.images = []
                        document.error = f"browser fallback stopped: {exc}"
                    except Exception as exc:
                        document.url = document.requested_url or document.url
                        document.error = f"browser fallback failed: {exc}"
            finally:
                await context.close()
                await browser.close()
    except Exception as exc:
        for document in pending:
            if not document.browser_rendered:
                document.error = f"browser fallback unavailable: {exc}"


async def gather_source_documents(
    candidates: Sequence[Any],
    validate_url: Callable[[str], Awaitable[None]],
    *,
    concurrency: int = 5,
    browser_fallback_enabled: bool = False,
    browser_timeout_seconds: float = 180,
    browser_settle_seconds: float = 5,
    browser_min_content_chars: int = 500,
    browser_site_cooldown_seconds: float = 30,
) -> list[SourceDocument]:
    """并发获取候选来源，同时按 URL 合并同源请求并维护两级缓存。

    内存锁保证同一进程内相同 URL 只下载一次；持久缓存跨行、跨任务复用正文。浏览器降级
    在普通 HTTP 获取完成后统一执行，最终成功文本再原子写入持久缓存。
    """

    semaphore = asyncio.Semaphore(concurrency)
    async with httpx.AsyncClient(follow_redirects=True, timeout=25) as client:

        async def guarded(candidate: Any, index: int) -> SourceDocument:
            normalized_url = normalize_source_url(candidate.source_url)
            metadata = getattr(candidate, "metadata", {})
            cache_scope = metadata.get("cache_scope") if isinstance(metadata, dict) else None
            # 动态榜单按任务快照隔离缓存：同一次十行复用一份正文，新任务即使在普通缓存
            # TTL 内也会重新读取当前榜单。scope 只参与内部缓存键，不发送给来源网站。
            cache_key = (
                f"{normalized_url}\ncache-scope:{cache_scope}"
                if cache_scope not in (None, "")
                else normalized_url
            )
            cached = _SOURCE_CACHE.get(cache_key)
            if cached is not None:
                return _cached_document(cached, candidate, index)
            lock = _SOURCE_CACHE_LOCKS.setdefault(cache_key, asyncio.Lock())
            async with lock:  # noqa: SIM117 - 分层显示进程内锁和跨进程锁的边界
                async with _cross_process_source_lock(cache_key):
                    cached = _SOURCE_CACHE.get(cache_key)
                    if cached is not None:
                        return _cached_document(cached, candidate, index)
                    persistent = _load_persistent_document(
                        cache_key,
                        candidate,
                        index,
                        normalized_url,
                    )
                    if persistent is not None:
                        if len(_SOURCE_CACHE) >= _SOURCE_CACHE_MAX_ITEMS:
                            _SOURCE_CACHE.pop(next(iter(_SOURCE_CACHE)))
                        _SOURCE_CACHE[cache_key] = copy.deepcopy(persistent)
                        return persistent
                    cooldown = _active_source_cooldown(cache_key, normalized_url)
                    if cooldown is not None:
                        return _cooldown_document(candidate, index, normalized_url, cache_key, cooldown)
                    await _reserve_host_request_slot(normalized_url)
                    async with semaphore:
                        normalized_candidate = copy.copy(candidate)
                        normalized_candidate.source_url = normalized_url
                        document = await fetch_source_document(
                            normalized_candidate,
                            index,
                            client,
                            validate_url,
                        )
                    document.requested_url = candidate.source_url
                    document.title = document.title or candidate.title
                    document.snippet = document.snippet or candidate.excerpt
                    document.normalized_url = normalized_url
                    document.cache_key = cache_key
                    if document.error:
                        _record_source_failure(cache_key, normalized_url, document)
                    elif document.text or document.images:
                        _clear_source_failure(cache_key, normalized_url)
                        if len(_SOURCE_CACHE) >= _SOURCE_CACHE_MAX_ITEMS:
                            _SOURCE_CACHE.pop(next(iter(_SOURCE_CACHE)))
                        _SOURCE_CACHE[cache_key] = copy.deepcopy(document)
                        _persist_document(cache_key, document)
                    return document

        documents = list(
            await asyncio.gather(
                *(guarded(candidate, index) for index, candidate in enumerate(candidates, start=1))
            )
        )
        if browser_fallback_enabled:
            await apply_browser_fallbacks(
                documents,
                client,
                validate_url,
                timeout_seconds=browser_timeout_seconds,
                settle_seconds=browser_settle_seconds,
                min_content_chars=browser_min_content_chars,
                site_cooldown_seconds=browser_site_cooldown_seconds,
            )
        for document in documents:
            cache_key = (
                document.cache_key
                or document.normalized_url
                or normalize_source_url(document.requested_url or document.url)
            )
            if not document.error and document.text:
                _clear_source_failure(cache_key, document.normalized_url or document.url)
                document.source_cooldown_until = None
                document.source_failure_category = None
                document.content_hash = (
                    document.content_hash or hashlib.sha256(document.text.encode()).hexdigest()
                )
                _SOURCE_CACHE[cache_key] = copy.deepcopy(document)
                _persist_document(cache_key, document)
            elif document.error and document.source_cooldown_until is None:
                _record_source_failure(
                    cache_key,
                    document.normalized_url or document.url,
                    document,
                )
        return documents


def build_contact_sheet(documents: Sequence[SourceDocument]) -> bytes | None:
    """按来源轮询选择图片拼成联系表，避免单一来源占满全部视觉槽位。"""

    images: list[ImageEvidence] = []
    for image_position in range(MAX_IMAGES_PER_SOURCE):
        for document in documents:
            if len(document.images) > image_position:
                images.append(document.images[image_position])
            if len(images) >= MAX_VISION_IMAGES:
                break
        if len(images) >= MAX_VISION_IMAGES:
            break
    if not images:
        return None
    cell_width, cell_height, label_height = 700, 520, 44
    columns = 2
    rows = (len(images) + columns - 1) // columns
    canvas = Image.new("RGB", (columns * cell_width, rows * cell_height), "white")
    draw = ImageDraw.Draw(canvas)
    for index, evidence in enumerate(images):
        with Image.open(io.BytesIO(evidence.png)) as image:
            image = image.convert("RGB")
            image.thumbnail((cell_width - 20, cell_height - label_height - 20))
            x = (index % columns) * cell_width + (cell_width - image.width) // 2
            y = (index // columns) * cell_height + label_height
            canvas.paste(image, (x, y))
        label = evidence.label[:95]
        label_position = (
            (index % columns) * cell_width + 10,
            (index // columns) * cell_height + 10,
        )
        draw.text(label_position, label, fill="black")
    output = io.BytesIO()
    canvas.save(output, format="PNG", optimize=True)
    return output.getvalue()
