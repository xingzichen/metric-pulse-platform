from __future__ import annotations

import io
import zipfile
from types import SimpleNamespace

import pymupdf
import pytest
from docx import Document
from openpyxl import Workbook
from PIL import Image

from metric_pulse.source_pipeline import (
    AttachmentReference,
    ImageEvidence,
    ImageReference,
    SourceDocument,
    browser_fallback_reason,
    build_contact_sheet,
    compact_json_records,
    extract_docx,
    extract_html_main_content,
    extract_pdf,
    extract_xlsx,
    fetch_source_document,
    gather_source_documents,
    looks_like_challenge_page,
    normalize_source_url,
    source_request_headers,
    validate_office_archive,
)


def test_html_extraction_keeps_main_content_and_drops_boilerplate() -> None:
    text, images, attachments = extract_html_main_content(
        b"""
        <html><body>
          <nav>unrelated navigation</nav>
          <main><h1>AI report</h1><p>Sweden created 42 AI companies in 2025.</p>
          <img src="/chart.png" alt="AI companies chart" width="800" height="500">
          <a href="/appendix.pdf">Download appendix</a>
          <img src="/logo.png" alt="logo" width="100" height="50"></main>
          <footer>unrelated footer</footer>
        </body></html>
        """,
        "https://example.com/report",
    )

    assert "42 AI companies" in text
    assert "unrelated navigation" not in text
    assert "unrelated footer" not in text
    assert "[[METRIC_PULSE_IMAGE:1]]" in text
    assert images == [
        ImageReference(
            url="https://example.com/chart.png",
            description=("alt: AI companies chart | 邻近正文: Sweden created 42 AI companies in 2025."),
            marker="[[METRIC_PULSE_IMAGE:1]]",
            referer="https://example.com/report",
        )
    ]
    assert attachments == [
        AttachmentReference(
            url="https://example.com/appendix.pdf",
            parent_url="https://example.com/report",
            anchor_text="Download appendix",
            surrounding_text="Download appendix",
            filename="appendix.pdf",
        )
    ]


def test_attachment_discovery_is_limited_to_main_and_marks_unsupported_formats() -> None:
    _, _, attachments = extract_html_main_content(
        b"""
        <html><body>
          <nav><a href="/nav.pdf">navigation download</a></nav>
          <main>
            <p>Official tables <a href="files/data.xlsx">data workbook</a></p>
            <p><a href="/slides.pptx">download slides</a></p>
            <a href="/about">ordinary page</a>
          </main>
          <footer><a href="/footer.zip">footer archive</a></footer>
        </body></html>
        """,
        "https://example.com/reports/2025",
    )

    assert [item.url for item in attachments] == [
        "https://example.com/reports/files/data.xlsx",
        "https://example.com/slides.pptx",
    ]
    assert attachments[0].supported is True
    assert attachments[1].supported is False
    assert attachments[1].unsupported_reason == "unsupported attachment format: .pptx"


def test_world_bank_page_is_normalized_to_official_structured_api() -> None:
    assert normalize_source_url("https://data.worldbank.org.cn/indicator/GB.XPD.RSDV.GD.ZS?locations=US") == (
        "https://api.worldbank.org/v2/zh/country/US/indicator/GB.XPD.RSDV.GD.ZS?format=json&per_page=20000"
    )


def test_nested_json_record_array_is_compacted_to_matchable_csv() -> None:
    text = compact_json_records(
        b'[{"page":1},[{"country":{"id":"US","value":"United States"},'
        b'"date":"2019","value":3.14297,"decimal":2}]]'
    )

    assert text is not None
    assert text.splitlines()[0] == "country,country_id,date,value,decimal"
    assert "United States,US,2019,3.14297,2" in text


def test_github_api_keeps_raw_json_for_dedicated_ranking_parser(monkeypatch) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    payload = b'{"total_count":1,"incomplete_results":false,"items":[{"full_name":"a/b"}]}'

    async def fake_download(_client, url, _validate_url, **_kwargs):
        return url, "application/json", payload

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "_download", fake_download)
    candidate = SimpleNamespace(
        source_url="https://api.github.com/search/repositories?q=stars",
        title="Ranking",
        excerpt=None,
    )

    document = asyncio.run(fetch_source_document(candidate, 1, object(), allow))

    assert document.text.startswith('{"total_count":1')
    assert '"items"' in document.text


def test_gather_expands_main_attachment_as_independent_source(monkeypatch, tmp_path) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    settings = pipeline.get_settings()
    monkeypatch.setattr(settings, "source_cache_root", tmp_path)
    monkeypatch.setattr(settings, "source_host_min_interval_seconds", 0)
    requests: list[tuple[str, str | None]] = []

    async def fake_download(_client, url, _validate_url, **kwargs):
        requests.append((url, kwargs.get("referer")))
        if url == "https://example.com/report":
            return (
                url,
                "text/html",
                b'<main><h1>Annual report</h1><p><a href="/data.csv">Download data</a></p></main>',
            )
        return url, "text/csv", b"country,value\nSweden,456"

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "_download", fake_download)
    candidate = SimpleNamespace(
        source_url="https://example.com/report",
        title="Report",
        excerpt=None,
        metadata={},
    )

    documents = asyncio.run(
        gather_source_documents(
            [candidate],
            allow,
            browser_fallback_enabled=False,
            attachment_max_per_parent=5,
            attachment_max_per_unit=8,
        )
    )

    assert len(documents) == 2
    assert documents[0].relation_type == "PRIMARY"
    assert documents[1].relation_type == "ATTACHMENT"
    assert documents[1].parent_url == "https://example.com/report"
    assert documents[1].attachment_filename == "data.csv"
    assert "Sweden,456" in documents[1].text
    assert requests[-1] == ("https://example.com/data.csv", "https://example.com/report")


def test_attachment_failure_isolated_from_usable_parent(monkeypatch, tmp_path) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    settings = pipeline.get_settings()
    monkeypatch.setattr(settings, "source_cache_root", tmp_path)
    monkeypatch.setattr(settings, "source_host_min_interval_seconds", 0)

    async def fake_download(_client, url, _validate_url, **_kwargs):
        assert url == "https://example.com/report"
        return (
            url,
            "text/html",
            b'<main><p>Usable parent value 42.</p><a href="/archive.zip">Download archive</a></main>',
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "_download", fake_download)
    candidate = SimpleNamespace(
        source_url="https://example.com/report",
        title="Report",
        excerpt=None,
        metadata={},
    )

    documents = asyncio.run(
        gather_source_documents([candidate], allow, browser_fallback_enabled=False)
    )

    assert "Usable parent value 42." in documents[0].text
    assert documents[1].relation_type == "ATTACHMENT"
    assert documents[1].error == "unsupported attachment format: .zip"


def test_attachment_limits_are_audited_and_nested_links_are_not_followed(monkeypatch, tmp_path) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    settings = pipeline.get_settings()
    monkeypatch.setattr(settings, "source_cache_root", tmp_path)
    monkeypatch.setattr(settings, "source_host_min_interval_seconds", 0)
    requested: list[str] = []

    async def fake_download(_client, url, _validate_url, **_kwargs):
        requested.append(url)
        if url == "https://example.com/report":
            return (
                url,
                "text/html",
                b"""<main><p>Parent</p>
                <a href="/first">Download first</a>
                <a href="/second.pdf">Download second</a></main>""",
            )
        assert url == "https://example.com/first"
        return (
            url,
            "text/html",
            b'<main><p>Nested attachment body 123</p><a href="/nested.pdf">Download nested</a></main>',
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "_download", fake_download)
    candidate = SimpleNamespace(
        source_url="https://example.com/report",
        title="Report",
        excerpt=None,
        metadata={},
    )

    documents = asyncio.run(
        gather_source_documents(
            [candidate],
            allow,
            browser_fallback_enabled=False,
            attachment_max_per_parent=1,
            attachment_max_per_unit=8,
        )
    )

    assert requested == ["https://example.com/report", "https://example.com/first"]
    assert len(documents) == 3
    assert "Nested attachment body 123" in documents[1].text
    assert documents[2].error == "ATTACHMENT_PARENT_LIMIT_EXCEEDED"
    assert all(document.url != "https://example.com/nested.pdf" for document in documents)


def test_pdf_extraction_includes_text_and_rendered_page() -> None:
    document = pymupdf.open()
    page = document.new_page()
    page.insert_text((72, 72), "AI investment reached 123 million in 2025")
    data = document.tobytes()
    document.close()

    text, images = extract_pdf(data, source_index=2)

    assert "123 million" in text
    assert images
    assert images[0].source_index == 2
    assert images[0].png.startswith(b"\x89PNG")


def test_extensionless_pdf_uses_content_type_instead_of_html_branch(monkeypatch) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pdf = pymupdf.open()
    pdf.new_page().insert_text((72, 72), "Attachment value 789")
    payload = pdf.tobytes()
    pdf.close()

    async def fake_download(_client, url, _validate_url, **_kwargs):
        return url, "application/pdf", payload, "official-report.pdf"

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "_download", fake_download)
    candidate = SimpleNamespace(
        source_url="https://example.com/download?id=42",
        title="Official report",
        excerpt=None,
        metadata={"relation_type": "ATTACHMENT"},
    )

    document = asyncio.run(fetch_source_document(candidate, 2, object(), allow))

    assert document.media_type == "application/pdf"
    assert document.attachment_filename == "official-report.pdf"
    assert "789" in document.text


def test_docx_extraction_includes_paragraphs_and_tables() -> None:
    document = Document()
    document.add_paragraph("AI talent report")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Country"
    table.cell(0, 1).text = "People"
    table.cell(1, 0).text = "Sweden"
    table.cell(1, 1).text = "456"
    buffer = io.BytesIO()
    document.save(buffer)

    text, _ = extract_docx(buffer.getvalue(), source_index=3)

    assert "AI talent report" in text
    assert "Sweden\t456" in text


def test_xlsx_extraction_includes_sheet_names_and_formula_results() -> None:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Indicators"
    sheet.append(["Country", "Value"])
    sheet.append(["Sweden", 456])
    buffer = io.BytesIO()
    workbook.save(buffer)

    text = extract_xlsx(buffer.getvalue())

    assert "Sheet: Indicators" in text
    assert "Country\tValue" in text
    assert "Sweden\t456" in text


def test_office_attachment_rejects_unsafe_compression_ratio() -> None:
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr("xl/worksheets/sheet1.xml", b"0" * 2_000_000)

    with pytest.raises(ValueError, match="compression ratio is unsafe"):
        validate_office_archive(buffer.getvalue())


def test_http_request_profile_is_browser_coherent_without_forged_client_hints() -> None:
    headers = source_request_headers(
        "https://cdn.example.net/report.pdf",
        referer="https://example.com/report?access_token=secret#download",
    )

    assert headers["User-Agent"].startswith("Mozilla/5.0")
    assert headers["Accept-Language"].startswith("zh-CN")
    assert headers["Referer"] == "https://example.com/"
    assert not any(key.lower().startswith(("sec-fetch", "sec-ch-ua")) for key in headers)


def test_contact_sheet_combines_visual_evidence() -> None:
    image = Image.new("RGB", (320, 200), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    document = SourceDocument(
        index=1,
        url="https://example.com/report",
        images=[ImageEvidence("Source 1 chart", buffer.getvalue(), 1)],
    )

    sheet = build_contact_sheet([document])

    assert sheet is not None
    assert sheet.startswith(b"\x89PNG")


def test_browser_fallback_targets_blocked_or_javascript_thin_pages() -> None:
    blocked = SourceDocument(
        index=1,
        url="https://example.com/report",
        http_status=403,
        error="HTTP 403: Forbidden",
    )
    thin = SourceDocument(
        index=2,
        url="https://example.com/dashboard",
        media_type="text/html",
        text="Loading...",
    )
    binary = SourceDocument(
        index=3,
        url="https://example.com/report.pdf",
        http_status=403,
        error="HTTP 403: Forbidden",
    )
    github_api = SourceDocument(
        index=4,
        url="https://api.github.com/search/repositories?q=stars",
        http_status=429,
        error="HTTP 429: Too Many Requests",
    )
    security_rejection = SourceDocument(
        index=5,
        url="https://example.com/private-redirect",
        error="Private, loopback, and reserved evidence addresses are not allowed",
    )

    assert browser_fallback_reason(blocked, min_content_chars=500) == "HTTP 403"
    assert "shorter than 500" in (browser_fallback_reason(thin, min_content_chars=500) or "")
    assert browser_fallback_reason(binary, min_content_chars=500) is None
    assert browser_fallback_reason(github_api, min_content_chars=500) is None
    assert browser_fallback_reason(security_rejection, min_content_chars=500) is None


def test_challenge_detection_does_not_flag_long_normal_article() -> None:
    assert looks_like_challenge_page("安全验证\n请输入验证码")
    assert looks_like_challenge_page('<script src="/cf-chl-/challenge-platform.js"></script>')
    long_article = "This research article discusses captcha systems. " + ("evidence " * 1_000)
    assert not looks_like_challenge_page(long_article)


def test_github_blob_url_is_normalized_to_raw_and_tracking_is_removed() -> None:
    normalized = normalize_source_url(
        "https://github.com/github/innovationgraph/blob/main/data/developers.csv?utm_source=test#readme"
    )

    assert normalized == ("https://raw.githubusercontent.com/github/innovationgraph/main/data/developers.csv")


def test_source_document_cache_survives_memory_cache_clear(monkeypatch, tmp_path) -> None:
    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    monkeypatch.setattr(pipeline.get_settings(), "source_cache_root", tmp_path)
    fetches = 0

    async def fake_fetch(candidate, index, client, validate_url):
        nonlocal fetches
        fetches += 1
        return SourceDocument(
            index=index,
            url=candidate.source_url,
            requested_url=candidate.source_url,
            normalized_url=candidate.source_url,
            media_type="text/csv",
            text="developers,iso2_code,year,quarter\n125033,EC,2021,3",
            content_hash="b" * 64,
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "fetch_source_document", fake_fetch)
    candidate = SimpleNamespace(
        source_url="https://example.com/developers.csv",
        title="Developers",
        excerpt=None,
    )

    first = __import__("asyncio").run(
        gather_source_documents([candidate], allow, browser_fallback_enabled=False)
    )
    pipeline._SOURCE_CACHE.clear()
    second = __import__("asyncio").run(
        gather_source_documents([candidate], allow, browser_fallback_enabled=False)
    )

    assert fetches == 1
    assert first[0].persistent_cache_hit is False
    assert second[0].persistent_cache_hit is True
    assert second[0].text == first[0].text


def test_persistent_source_cache_restores_image_bytes_without_refetch(monkeypatch, tmp_path) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    settings = pipeline.get_settings()
    monkeypatch.setattr(settings, "source_cache_root", tmp_path)
    monkeypatch.setattr(settings, "source_host_min_interval_seconds", 0)
    fetches = 0
    image = Image.new("RGB", (24, 24), "white")
    buffer = io.BytesIO()
    image.save(buffer, format="PNG")
    png = buffer.getvalue()

    async def fake_fetch(candidate, index, client, validate_url):
        nonlocal fetches
        fetches += 1
        return SourceDocument(
            index=index,
            url=candidate.source_url,
            media_type="text/html",
            text="Article body [[METRIC_PULSE_IMAGE:1]]",
            images=[
                ImageEvidence(
                    "Source 1 chart",
                    png,
                    index,
                    description="Figure 1: score table",
                    marker="[[METRIC_PULSE_IMAGE:1]]",
                )
            ],
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "fetch_source_document", fake_fetch)
    candidate = SimpleNamespace(
        source_url="https://example.com/image-article",
        title="Image article",
        excerpt=None,
    )

    first = asyncio.run(gather_source_documents([candidate], allow, browser_fallback_enabled=False))
    pipeline._SOURCE_CACHE.clear()
    second = asyncio.run(gather_source_documents([candidate], allow, browser_fallback_enabled=False))

    assert fetches == 1
    assert first[0].images[0].png == png
    assert second[0].persistent_cache_hit is True
    assert second[0].images[0].png == png
    assert second[0].images[0].description == "Figure 1: score table"


def test_challenge_negative_cache_blocks_repeated_url_and_same_host(monkeypatch, tmp_path) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    settings = pipeline.get_settings()
    monkeypatch.setattr(settings, "source_cache_root", tmp_path)
    monkeypatch.setattr(settings, "source_host_min_interval_seconds", 0)
    monkeypatch.setattr(settings, "source_challenge_cooldown_seconds", 600)
    monkeypatch.setattr(settings, "source_cooldown_max_seconds", 600)
    fetches = 0

    async def fake_fetch(candidate, index, client, validate_url):
        nonlocal fetches
        fetches += 1
        return SourceDocument(
            index=index,
            url=candidate.source_url,
            http_status=403,
            error="HTTP 403: security challenge captcha",
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "fetch_source_document", fake_fetch)

    def candidate(path: str):
        return SimpleNamespace(
            source_url=f"https://blocked.example.com/{path}",
            title="Blocked article",
            excerpt=None,
        )

    first = asyncio.run(
        gather_source_documents([candidate("one")], allow, browser_fallback_enabled=False)
    )
    repeated = asyncio.run(
        gather_source_documents([candidate("one")], allow, browser_fallback_enabled=False)
    )
    same_host = asyncio.run(
        gather_source_documents([candidate("two")], allow, browser_fallback_enabled=False)
    )

    assert fetches == 1
    assert first[0].source_failure_category == "CHALLENGE"
    assert repeated[0].error == "source cooldown active: CHALLENGE"
    assert same_host[0].error == "source cooldown active: CHALLENGE"
    assert repeated[0].retry_after_seconds > 0
    normalized = pipeline.normalize_source_url(candidate("one").source_url)
    pipeline._clear_source_failure(normalized, normalized)
    assert pipeline._active_source_cooldown(normalized, normalized) is None


def test_security_policy_rejection_is_not_negative_cached(monkeypatch, tmp_path) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    settings = pipeline.get_settings()
    monkeypatch.setattr(settings, "source_cache_root", tmp_path)
    monkeypatch.setattr(settings, "source_host_min_interval_seconds", 0)
    fetches = 0

    async def fake_fetch(candidate, index, client, validate_url):
        nonlocal fetches
        fetches += 1
        return SourceDocument(
            index=index,
            url=candidate.source_url,
            error="Private, loopback, and reserved evidence addresses are not allowed",
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "fetch_source_document", fake_fetch)
    candidate = SimpleNamespace(
        source_url="https://example.com/private-redirect",
        title="Invalid direct source",
        excerpt=None,
    )

    first = asyncio.run(gather_source_documents([candidate], allow, browser_fallback_enabled=True))
    second = asyncio.run(gather_source_documents([candidate], allow, browser_fallback_enabled=True))

    assert fetches == 2
    assert first[0].source_cooldown_until is None
    assert second[0].source_cooldown_until is None
    normalized = pipeline.normalize_source_url(candidate.source_url)
    assert pipeline._active_source_cooldown(normalized, normalized) is None


def test_legacy_security_policy_cooldown_is_ignored(monkeypatch, tmp_path) -> None:
    import time

    import metric_pulse.source_pipeline as pipeline

    monkeypatch.setattr(pipeline.get_settings(), "source_cache_root", tmp_path)
    normalized = "https://example.com/private-redirect"
    path = pipeline._source_failure_path(normalized)
    pipeline._write_json_file(
        path,
        {
            "normalized_url": normalized,
            "hostname": "example.com",
            "category": "TRANSIENT",
            "failure_count": 7,
            "blocked_until": time.time() + 3_600,
            "error": "Private, loopback, and reserved evidence addresses are not allowed",
        },
    )

    assert pipeline._active_source_cooldown(normalized, normalized) is None


def test_dynamic_snapshot_cache_is_reused_within_scope_but_refetched_for_new_scope(
    monkeypatch,
    tmp_path,
) -> None:
    import asyncio

    import metric_pulse.source_pipeline as pipeline

    pipeline._SOURCE_CACHE.clear()
    pipeline._SOURCE_CACHE_LOCKS.clear()
    monkeypatch.setattr(pipeline.get_settings(), "source_cache_root", tmp_path)
    fetches = 0

    async def fake_fetch(candidate, index, client, validate_url):
        nonlocal fetches
        fetches += 1
        return SourceDocument(
            index=index,
            url=candidate.source_url,
            requested_url=candidate.source_url,
            normalized_url=candidate.source_url,
            media_type="application/json",
            text=f'{{"fetch": {fetches}}}',
        )

    async def allow(_url):
        return None

    monkeypatch.setattr(pipeline, "fetch_source_document", fake_fetch)

    def candidate(scope: str):
        return SimpleNamespace(
            source_url="https://api.github.com/search/repositories?q=stars",
            title="Ranking",
            excerpt=None,
            metadata={"cache_scope": scope},
        )

    first = asyncio.run(gather_source_documents([candidate("snapshot-a")], allow))
    same_snapshot = asyncio.run(gather_source_documents([candidate("snapshot-a")], allow))
    new_snapshot = asyncio.run(gather_source_documents([candidate("snapshot-b")], allow))

    assert fetches == 2
    assert same_snapshot[0].cache_hit is True
    assert first[0].text == same_snapshot[0].text
    assert new_snapshot[0].text != first[0].text
    assert all(
        document.normalized_url == "https://api.github.com/search/repositories?q=stars"
        for document in (first[0], same_snapshot[0], new_snapshot[0])
    )
